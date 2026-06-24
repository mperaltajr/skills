"""
emit_dot_dash.py — Project a narrative brief into a dot-dash storyline file
==========================================================================

Reads a `narrative-brief-<topic>.md` and writes companion
`dot-dash-<topic>.docx`, `dot-dash-<topic>.md`, and `dot-dash-<topic>.html`
files that lay out the deck's argument in McKinsey-style dot-dash format:

    • Slide governing thought (the action title)
      – Evidence sentence as prose (no labels, no ALL-CAPS prefixes)
      – Evidence sentence as prose
      – Takeaway – the so-what synthesizing the dashes above

Read the dot headlines + Takeaway lines top-to-bottom and you should get
the deck's argument as a coherent story. If they don't, the storyline is
broken — that's the value of generating this file before any slide is
built (the read-down test, Minto Pyramid Principle).

The dot-dash is the human-facing artifact. Its job is to read like a
consultant wrote it, not to carry the structured `**HEADING** — body`
shape the slide-builder translator needs. Bullets in the brief that use
that shape are stripped to prose here. The brief retains the structure
for slide-builder downstream — two views of the same content.

Usage:
    py -3 emit_dot_dash.py "<path to narrative-brief-*.md>"

Output:
- If the brief lives in `<root>/_session/narrative-brief-<topic>.md`
  → writes `<root>/dot-dash-<topic>.docx` + .md + .html
- Otherwise → writes all three next to the brief

The script never modifies the brief.

Public API (callable from review_html.py):
  parse_brief_for_dot_dash(brief_text) -> {header, slides}
    Structured data, no rendering.
  render_dot_dash(brief_text) -> str
    Legacy markdown render (kept for back-compat).
  render_dot_dash_md(data) -> str
    Markdown from parsed data.
  render_dot_dash_html(data, *, standalone=True) -> str
    HTML render. standalone=True returns a full <html> page; False returns
    an inner fragment suitable for embedding in another HTML page.
  render_dot_dash_docx(data, output_path) -> None
    Word document render — Heading 1 deck title, prose dashes, bold
    Takeaway label. Matches the hand-written house-style example.
"""

from __future__ import annotations

import argparse
import html
import pathlib
import re
import sys
from typing import Any, Dict, List


# ---------------------------------------------------------------------------
# Field extraction helpers
# ---------------------------------------------------------------------------

def _extract_field(block: str, label: str) -> str:
    """Pull the value of a `**Label:**` field out of a block.

    The field value runs from the colon to the next `**` heading or blank
    line followed by another heading. Returns "" if not found.
    """
    pattern = rf"\*\*{re.escape(label)}:\*\*\s*(.+?)(?=\n\s*\*\*|\n---|\Z)"
    m = re.search(pattern, block, re.DOTALL)
    return m.group(1).strip() if m else ""


# Brief-format prefix that needs stripping when projected into the dot-dash.
# The brief carries `- **HEADING** — body sentence` because slide-builder's
# translator parses that shape into structured card/pillar overrides. For
# the dot-dash artifact (human-facing) we want prose, so we drop the bold
# heading and its trailing separator. Handles both em-dash (—) and en-dash
# (–) and plain hyphen, with optional whitespace.
_LABEL_PREFIX_RE = re.compile(
    r"^\*\*[^*]+\*\*\s*[—–\-]\s*"
)

# Key-value structured bullets (`**Title:** Slide Labs`, `**Tagline:** ...`)
# show up on cover/transition/closing slides. They are slide-construction
# data, not story content. Detect so the renderer can elide them from the
# dot-dash narrative.
_KV_BULLET_RE = re.compile(r"^\*\*[^*]+:\*\*\s*")


def _strip_bullet_label(text: str) -> str:
    """Strip a `**HEADING** — ` prefix from a bullet so it reads as prose.

    Returns the bullet unchanged when no recognized prefix is present.
    Capitalizes the first letter of the resulting prose (sentence case)
    when the strip leaves a lowercase opener — labels in the brief are
    often followed by a sentence fragment that, freed of the label, no
    longer starts cleanly.
    """
    stripped = _LABEL_PREFIX_RE.sub("", text, count=1)
    if stripped != text and stripped and stripped[0].islower():
        stripped = stripped[0].upper() + stripped[1:]
    return stripped


def _is_kv_bullet(text: str) -> bool:
    """True when the bullet is a `**Field:** value` structured pair — these
    appear on cover/transition/closing slides and are not prose content."""
    return bool(_KV_BULLET_RE.match(text))


def _extract_evidence_bullets(block: str) -> list[str]:
    """Pull evidence bullets out of the `**Evidence / content:**` field.

    Recognises three line shapes:
      - Top-level bullets (`- ` / `* `)            → kept as dashes (label-stripped)
      - Indented continuation bullets               → merged into previous dash
      - Non-bullet label lines ending with `:`     → buffered, used to label
                                                     any table that follows
      - Table rows (`|...|`)                        → collapsed into one dash,
                                                     anchored to the buffered
                                                     label OR (if absent) to
                                                     the prior bullet

    Bullets that are key-value pairs (`**Title:** ...`) are returned as-is —
    callers (cover/transition slides) decide whether to render them. All
    other bullets have their `**HEADING** — ` prefix stripped so the dot-
    dash reads as prose (not labeled callouts).

    Stray prose lines that are neither bullets nor labels are ignored.
    """
    evidence = _extract_field(block, "Evidence / content")
    if not evidence:
        return []

    bullets: list[str] = []
    pending_label: str | None = None  # non-bullet "X:" line awaiting context
    in_table = False

    for raw_line in evidence.splitlines():
        line = raw_line.rstrip()
        stripped = line.lstrip()

        if not stripped:
            continue

        # Table row — emit one summary dash per table, then skip subsequent rows
        if stripped.startswith("|"):
            if not in_table:
                in_table = True
                if pending_label:
                    bullets.append(f"{pending_label.rstrip(':')} table (see brief)")
                    pending_label = None
                elif bullets:
                    last = bullets[-1].rstrip(":. ")
                    if "table" in last.lower():
                        bullets[-1] = last  # strip dangling punctuation only
                    else:
                        bullets[-1] = last + " (see table in brief)"
                else:
                    bullets.append("(table in brief)")
            continue
        else:
            in_table = False

        # Top-level bullets
        if line.startswith(("- ", "* ")):
            raw = line[2:].strip()
            # Key-value bullets (`**Title:** ...`, `**Primary ask:** ...`)
            # are slide-construction data — title, presenter, ask labels.
            # They never carry narrative content, so drop them entirely
            # from the dot-dash. The dot + Takeaway carry the story.
            if _is_kv_bullet(raw):
                pending_label = None
                continue
            bullets.append(_strip_bullet_label(raw))
            pending_label = None  # bullets supersede any pending label
            continue

        # Indented continuation
        if line.startswith(("  - ", "  * ")):
            if bullets:
                bullets[-1] = bullets[-1] + " — " + line.lstrip()[2:].strip()
            continue

        # Non-bullet line — track as potential label if it ends with ":"
        if line.endswith(":"):
            pending_label = line
        # else: stray prose, ignore

    return bullets


def _exhibit_line(block: str) -> str | None:
    """Build an `Exhibit:` dash from Chart type + Chart data if present."""
    chart_type = _extract_field(block, "Chart type")
    if not chart_type or chart_type.lower().strip() in {"none", "n/a", ""}:
        return None
    chart_data = _extract_field(block, "Chart data")
    # Trim chart_data to a one-line summary if it's a table or path
    chart_summary = ""
    if chart_data:
        first_line = chart_data.splitlines()[0].strip()
        if first_line.startswith("|"):
            chart_summary = "(data table in brief)"
        elif first_line.lower().startswith(("tbd", "placeholder")):
            chart_summary = "(data TBD)"
        else:
            chart_summary = first_line[:120]
    return f"Exhibit: {chart_type} chart — {chart_summary}".rstrip(" —")


# ---------------------------------------------------------------------------
# Per-slide and deck-level parsers
# ---------------------------------------------------------------------------

SLIDE_HEADER_RE = re.compile(r"^###\s+Slide\s+\d+", re.MULTILINE)


def _split_slides(brief_text: str) -> list[str]:
    """Split the brief into per-slide blocks based on `### Slide N` headers."""
    # Find header positions, slice from each header to the next (or end)
    matches = list(SLIDE_HEADER_RE.finditer(brief_text))
    if not matches:
        return []
    blocks: list[str] = []
    for i, m in enumerate(matches):
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(brief_text)
        blocks.append(brief_text[start:end])
    return blocks


def _slide_title(block: str) -> str:
    """Return the slide header text (e.g. 'Slide 1 — Overall Program Status')."""
    first_line = block.splitlines()[0]
    return first_line.lstrip("# ").strip()


def _parse_deck_header(brief_text: str) -> dict[str, str]:
    """Pull the deck-level fields shown at the top of the dot-dash file."""
    # Truncate to just the pre-Sequence section
    sequence_idx = brief_text.find("## Sequence")
    head = brief_text[:sequence_idx] if sequence_idx > 0 else brief_text

    def _grab(heading: str) -> str:
        m = re.search(
            rf"##\s+{re.escape(heading)}\s*\n+([^\n#]+(?:\n(?!##)[^\n]*)*)",
            head,
        )
        return m.group(1).strip() if m else ""

    return {
        "topic": _grab_title(brief_text),
        "deck_type": _grab("Deck type"),
        "governing": _grab("Governing thought (the whole deck)"),
        "audience": _grab("Audience"),
        "belief_break": _grab_inline(head, "Audience assumption to break"),
        "belief_leave": _grab_inline(head, "Audience belief to leave with"),
        "say_back": _grab_inline(head, "The single sentence the room should say back"),
    }


def _grab_title(brief_text: str) -> str:
    m = re.search(r"^#\s+Narrative brief:\s*(.+)$", brief_text, re.MULTILINE)
    return m.group(1).strip() if m else ""


def _grab_inline(text: str, label: str) -> str:
    m = re.search(rf"\*\*{re.escape(label)}:\*\*\s*([^\n]+)", text)
    return m.group(1).strip() if m else ""


# ---------------------------------------------------------------------------
# Parse — structured intermediate form (used by both md and html renderers)
# ---------------------------------------------------------------------------

_NO_NARRATIVE_ARCHETYPES = {
    "cover / title", "cover", "title",
    # Transition cards that introduce a section also carry no story content.
    "transition", "section divider", "section / divider",
}


def _is_no_narrative_slide(archetype: str, governing: str, bullets: list[str]) -> bool:
    """True when a slide carries no narrative dashes — typically a cover,
    transition, or closing-CTA slide whose bullets are key-value pairs.

    These slides get their dashes elided from the dot-dash so the artifact
    reads as a continuous argument rather than tripping over title/tagline
    metadata that belongs on a slide, not in a story.
    """
    archetype_n = (archetype or "").strip().lower()
    if archetype_n in _NO_NARRATIVE_ARCHETYPES:
        return True
    # Heuristic fallback: governing thought missing AND bullets are
    # predominantly structured key-value pairs.
    governing_missing = (
        not governing
        or "no governing thought" in governing.lower()
        or "transition card" in governing.lower()
    )
    if governing_missing and bullets:
        kv_count = sum(1 for b in bullets if _is_kv_bullet(b))
        if kv_count / len(bullets) > 0.5:
            return True
    return False


def parse_brief_for_dot_dash(brief_text: str) -> Dict[str, Any]:
    """Parse a narrative brief into structured dot-dash data.

    Returns:
      {
        "header": {topic, deck_type, governing, audience, belief_break,
                   belief_leave, say_back},
        "slides": [
          {"title": "Slide 2 — The Problem",
           "archetype": "Context / Situation",
           "governing": "...",
           "so_what": "...",         # used as the Takeaway dash
           "bullets": ["...", "..."],  # prose, labels stripped
           "exhibit": "Exhibit: bar chart — ...",
           "no_narrative": False},    # True for covers / transitions
          ...
        ]
      }
    """
    header = _parse_deck_header(brief_text)
    blocks = _split_slides(brief_text)
    slides: List[Dict[str, Any]] = []
    for block in blocks:
        archetype = _extract_field(block, "Archetype")
        # The brief stores so-what under either header form depending on
        # the slide kind. Try the canonical first, then the cover-shortened.
        so_what = (
            _extract_field(block, "So-what (the takeaway)")
            or _extract_field(block, "So-what")
        )
        # Cover/transition placeholder text ("[Cover slide — no so-what required]")
        # shouldn't surface as a Takeaway.
        if so_what.startswith("[") and so_what.endswith("]"):
            so_what = ""
        # Governing thought has the same conditional shape.
        governing = (
            _extract_field(block, "Governing thought (the claim)")
            or _extract_field(block, "Governing thought")
        )
        if governing.startswith("[") and governing.endswith("]"):
            governing = ""
        bullets = _extract_evidence_bullets(block)
        slides.append({
            "title": _slide_title(block),
            "archetype": archetype,
            "governing": governing,
            "so_what": so_what,
            "bullets": bullets,
            "exhibit": _exhibit_line(block),
            "no_narrative": _is_no_narrative_slide(archetype, governing, bullets),
        })
    return {"header": header, "slides": slides}


# ---------------------------------------------------------------------------
# Render — markdown
# ---------------------------------------------------------------------------

def render_dot_dash_md(data: Dict[str, Any]) -> str:
    """Render the structured dot-dash data as the canonical markdown form."""
    header = data["header"]
    slides = data["slides"]

    lines: list[str] = []
    lines.append(f"# Dot-dash storyline: {header['topic']}")
    lines.append("")
    if header["deck_type"]:
        lines.append(f"**Deck type:** {header['deck_type']}")
    if header["governing"]:
        lines.append(f"**Governing thought (whole deck):** {header['governing']}")
    if header["audience"]:
        audience_first_line = header["audience"].splitlines()[0]
        lines.append(f"**Audience:** {audience_first_line}")
    if header["belief_break"]:
        lines.append(f"**Belief to break:** {header['belief_break']}")
    if header["belief_leave"]:
        lines.append(f"**Belief to leave with:** {header['belief_leave']}")
    if header["say_back"]:
        lines.append(f"**The room should say back:** {header['say_back']}")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("> Read the dots top-to-bottom — they should form the deck's argument as a single coherent story. If the dots-alone don't make sense, the storyline is broken.")
    lines.append("")
    lines.append("---")
    lines.append("")

    if not slides:
        lines.append("_No slides found in the brief._")
        return "\n".join(lines) + "\n"

    for s in slides:
        lines.append(f"_{s['title']}_")
        lines.append("")
        if s["governing"]:
            lines.append(f"**• {s['governing']}**")
        else:
            # Cover / transition placeholder text — no governing thought
            # by design. Mark explicitly so the read-down test ignores it.
            lines.append(f"**• _(cover / transition — no governing thought)_**")
        # Covers and transitions ship no narrative dashes — the title and
        # any key-value metadata live on the slide itself.
        if not s.get("no_narrative"):
            for b in s["bullets"]:
                lines.append(f"  – {b}")
            if s["exhibit"]:
                lines.append(f"  – {s['exhibit']}")
            if s.get("so_what"):
                # Takeaway is the final dash — synthesizes the dashes above
                # into the slide's so-what. En-dash separator matches the
                # hand-written house-style example.
                lines.append(f"  – **Takeaway –** {s['so_what']}")
        lines.append("")

    return "\n".join(lines) + "\n"


# ---------------------------------------------------------------------------
# Render — HTML (standalone page OR embeddable fragment)
# ---------------------------------------------------------------------------

_DOT_DASH_CSS = """
:root {
  --bg: #0F172A;
  --panel: #1E293B;
  --text: #E2E8F0;
  --text-dim: #94A3B8;
  --accent: #A100FF;
  --accent-soft: #C780FF;
  --border: #334155;
  --rule: #475569;
}
.dd-container { font-family: 'Inter', -apple-system, 'Segoe UI', sans-serif; line-height: 1.5; color: var(--text); }
.dd-title { font-size: 22px; font-weight: 800; color: var(--text); margin: 0 0 14px; letter-spacing: -0.2px; }
.dd-deck-meta { display: grid; grid-template-columns: max-content 1fr; gap: 4px 14px; font-size: 13px; margin-bottom: 16px; }
.dd-deck-meta .lbl { color: var(--text-dim); font-weight: 700; text-transform: uppercase; font-size: 10px; letter-spacing: 1px; padding-top: 3px; }
.dd-deck-meta .val { color: var(--text); }
.dd-deck-meta .val.gov { color: var(--accent-soft); font-weight: 600; }
.dd-callout { background: rgba(161,0,255,0.08); border-left: 3px solid var(--accent); padding: 10px 14px; font-size: 12px; color: var(--text-dim); font-style: italic; margin: 14px 0; }
.dd-slide { padding: 10px 0 14px; border-top: 1px solid var(--border); }
.dd-slide:first-of-type { border-top: 0; }
.dd-slide-title { font-size: 11px; font-weight: 700; color: var(--text-dim); text-transform: uppercase; letter-spacing: 1.5px; margin-bottom: 6px; }
.dd-gov { font-size: 15px; font-weight: 700; color: var(--text); margin: 0 0 8px; line-height: 1.35; }
.dd-gov::before { content: "● "; color: var(--accent); }
.dd-gov.missing { color: var(--text-dim); font-style: italic; font-weight: 500; }
.dd-bullets { list-style: none; margin: 0; padding-left: 22px; }
.dd-bullets li { font-size: 13px; color: var(--text-dim); padding: 2px 0; position: relative; }
.dd-bullets li::before { content: "– "; color: var(--accent-soft); position: absolute; left: -16px; }
.dd-bullets li.exhibit { color: var(--accent-soft); font-style: italic; }
.dd-bullets li.takeaway { color: var(--text); padding-top: 6px; }
.dd-bullets li.takeaway strong { color: var(--accent-soft); margin-right: 4px; }
.dd-empty { color: var(--text-dim); font-style: italic; }
"""


def _h(text: str) -> str:
    return html.escape(text or "")


def render_dot_dash_html(data: Dict[str, Any], *, standalone: bool = True) -> str:
    """Render the dot-dash data as HTML.

    standalone=True  -> full <html><body> page with embedded CSS.
    standalone=False -> inner content only (a <div class="dd-container">),
                        for embedding in another HTML doc that supplies CSS.
    """
    header = data["header"]
    slides = data["slides"]

    parts: list[str] = []
    parts.append('<div class="dd-container">')
    parts.append(f'<h1 class="dd-title">Dot-dash storyline: {_h(header.get("topic", ""))}</h1>')

    meta_rows: list[tuple[str, str, str]] = []  # (label, value, css_class)
    if header.get("deck_type"):
        meta_rows.append(("Deck type", header["deck_type"], ""))
    if header.get("governing"):
        meta_rows.append(("Governing thought", header["governing"], "gov"))
    if header.get("audience"):
        meta_rows.append(("Audience", header["audience"].splitlines()[0], ""))
    if header.get("belief_break"):
        meta_rows.append(("Belief to break", header["belief_break"], ""))
    if header.get("belief_leave"):
        meta_rows.append(("Belief to leave with", header["belief_leave"], ""))
    if header.get("say_back"):
        meta_rows.append(("Room should say back", header["say_back"], ""))

    if meta_rows:
        parts.append('<div class="dd-deck-meta">')
        for label, value, css in meta_rows:
            cls = f"val {css}".strip()
            parts.append(f'<div class="lbl">{_h(label)}</div><div class="{cls}">{_h(value)}</div>')
        parts.append('</div>')

    parts.append(
        '<div class="dd-callout">Read the dots top-to-bottom — they should form the '
        "deck's argument as a single coherent story. If the dots-alone don't make sense, "
        'the storyline is broken.</div>'
    )

    if not slides:
        parts.append('<div class="dd-empty">No slides found in the brief.</div>')
    else:
        for s in slides:
            parts.append('<div class="dd-slide">')
            parts.append(f'<div class="dd-slide-title">{_h(s["title"])}</div>')
            if s["governing"]:
                parts.append(f'<div class="dd-gov">{_h(s["governing"])}</div>')
            else:
                parts.append('<div class="dd-gov missing">(cover / transition — no governing thought)</div>')
            if not s.get("no_narrative"):
                items: list[str] = []
                for b in s["bullets"]:
                    items.append(f'<li>{_h(b)}</li>')
                if s["exhibit"]:
                    items.append(f'<li class="exhibit">{_h(s["exhibit"])}</li>')
                if s.get("so_what"):
                    items.append(
                        f'<li class="takeaway"><strong>Takeaway –</strong> '
                        f'{_h(s["so_what"])}</li>'
                    )
                if items:
                    parts.append('<ul class="dd-bullets">')
                    parts.extend(items)
                    parts.append('</ul>')
            parts.append('</div>')

    parts.append('</div>')
    body = "\n".join(parts)

    if not standalone:
        return body

    # Standalone page: wrap with <html> + CSS. Fills the full viewport;
    # padding gives breathing room without forcing a fixed-width column.
    topic = _h(header.get("topic", "Dot-dash"))
    return (
        f"<!DOCTYPE html><html lang=\"en\"><head><meta charset=\"UTF-8\">"
        f"<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">"
        f"<title>Dot-dash: {topic}</title>"
        f"<style>html, body {{ margin:0; padding:0; }} "
        f"body {{ background:#0F172A; padding:40px 56px; min-height:100vh; box-sizing:border-box; }}\n"
        f"{_DOT_DASH_CSS}</style>"
        f"</head><body>{body}</body></html>"
    )


# ---------------------------------------------------------------------------
# Render — Word document (the canonical human-facing artifact)
# ---------------------------------------------------------------------------

def render_dot_dash_docx(data: Dict[str, Any], output_path: pathlib.Path) -> None:
    """Write the dot-dash data as a .docx file readable in Word.

    Layout matches the hand-written house-style reference:
      - Heading 1 deck title
      - Deck-level metadata as a key:value block (bold labels)
      - Per-slide block: bold dot headline (the governing thought) as a
        bulleted item, then indented dashes for evidence, then a bold
        "Takeaway –" dash synthesizing the so-what.

    Word is the canonical artifact because consultants edit, comment, and
    track-change in Word — markdown gets mangled in real review workflows.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    header = data["header"]
    slides = data["slides"]

    doc = Document()

    # Default style — Arial 11pt, the consulting-deliverable baseline.
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    # Title.
    topic = header.get("topic") or "Dot-dash storyline"
    doc.add_heading(f"Dot-dash storyline: {topic}", level=1)

    # Deck metadata block — bold label, value on the same line.
    def _meta(label: str, value: str) -> None:
        if not value:
            return
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label}: ")
        run_l.bold = True
        p.add_run(value)

    _meta("Deck type", header.get("deck_type", ""))
    _meta("Governing thought (whole deck)", header.get("governing", ""))
    if header.get("audience"):
        _meta("Audience", header["audience"].splitlines()[0])
    _meta("Belief to break", header.get("belief_break", ""))
    _meta("Belief to leave with", header.get("belief_leave", ""))
    _meta("The room should say back", header.get("say_back", ""))

    # Read-down callout — italic, lighter to set it apart from content.
    callout = doc.add_paragraph()
    r = callout.add_run(
        "Read the dot headlines top-to-bottom — together they should form "
        "the deck's argument as a coherent story. If they don't, the "
        "storyline is broken."
    )
    r.italic = True
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    if not slides:
        doc.add_paragraph("No slides found in the brief.")
        doc.save(str(output_path))
        return

    for s in slides:
        # Slide-title subtitle (e.g., "Slide 2 — The Problem").
        sub = doc.add_paragraph()
        sub.paragraph_format.space_before = Pt(14)
        sr = sub.add_run(s["title"])
        sr.italic = True
        sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        # The dot — bold governing thought, bulleted at level 0.
        dot_text = (
            s["governing"] if s["governing"]
            else "(cover / transition — no governing thought)"
        )
        p_dot = doc.add_paragraph(style="List Bullet")
        rd = p_dot.add_run(dot_text)
        rd.bold = True
        if not s["governing"]:
            rd.italic = True
            rd.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Dashes — covers/transitions get none.
        if s.get("no_narrative"):
            continue

        for b in s["bullets"]:
            p_dash = doc.add_paragraph(style="List Bullet 2")
            p_dash.paragraph_format.space_after = Pt(2)
            p_dash.add_run(b)

        if s.get("exhibit"):
            p_ex = doc.add_paragraph(style="List Bullet 2")
            p_ex.paragraph_format.space_after = Pt(2)
            er = p_ex.add_run(s["exhibit"])
            er.italic = True

        if s.get("so_what"):
            p_take = doc.add_paragraph(style="List Bullet 2")
            p_take.paragraph_format.space_before = Pt(4)
            tr_label = p_take.add_run("Takeaway – ")
            tr_label.bold = True
            p_take.add_run(s["so_what"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# Back-compat alias — legacy callers use render_dot_dash(brief_text) -> markdown.
def render_dot_dash(brief_text: str) -> str:
    """DEPRECATED — prefer parse_brief_for_dot_dash + render_dot_dash_md.
    Kept for any external caller that imported it directly.
    """
    return render_dot_dash_md(parse_brief_for_dot_dash(brief_text))


# Public CSS accessor — review_html.py uses this to inline the dot-dash styles.
def dot_dash_css() -> str:
    """Return the CSS rules used by render_dot_dash_html. Callers that embed
    the fragment via render_dot_dash_html(data, standalone=False) must include
    this CSS in their page's <style> block.
    """
    return _DOT_DASH_CSS


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _output_path(brief_path: pathlib.Path) -> pathlib.Path:
    """Determine where the dot-dash file should be written.

    Convention: brief lives in `<root>/_session/`, dot-dash goes to `<root>/`.
    If the brief is not inside a `_session/` folder, fall back to writing
    next to the brief.
    """
    topic_stem = brief_path.stem
    if topic_stem.startswith("narrative-brief-"):
        topic = topic_stem[len("narrative-brief-"):]
    else:
        topic = topic_stem
    out_name = f"dot-dash-{topic}.md"

    if brief_path.parent.name == "_session":
        return brief_path.parent.parent / out_name
    return brief_path.parent / out_name


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Project a narrative brief into a dot-dash storyline file."
    )
    parser.add_argument(
        "brief",
        help="Path to the narrative-brief-<topic>.md file",
    )
    args = parser.parse_args()

    brief_path = pathlib.Path(args.brief).resolve()
    if not brief_path.exists():
        print(f"ERROR: brief not found: {brief_path}", file=sys.stderr)
        return 1
    if not brief_path.is_file():
        print(f"ERROR: not a file: {brief_path}", file=sys.stderr)
        return 1

    brief_text = brief_path.read_text(encoding="utf-8")
    data = parse_brief_for_dot_dash(brief_text)

    md_path = _output_path(brief_path)
    html_path = md_path.with_suffix(".html")
    docx_path = md_path.with_suffix(".docx")

    md_path.write_text(render_dot_dash_md(data), encoding="utf-8")
    html_path.write_text(render_dot_dash_html(data, standalone=True), encoding="utf-8")
    try:
        render_dot_dash_docx(data, docx_path)
        docx_status = str(docx_path)
    except ImportError as exc:
        # python-docx not installed — emit a clear pointer rather than
        # crashing. .docx is the canonical artifact, but the .md companion
        # still ships.
        docx_status = (
            f"SKIPPED ({exc.__class__.__name__}: {exc}). "
            "Install python-docx: `py -3 -m pip install python-docx`"
        )

    print(f"Dot-dash written:")
    print(f"  {docx_path if 'SKIPPED' not in docx_status else docx_status}")
    print(f"  {md_path}")
    print(f"  {html_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
